from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.views.generic import ListView, DetailView, TemplateView
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.http import JsonResponse, HttpResponseRedirect, HttpResponse, Http404
from django.views.decorators.http import require_http_methods
from django.db import models, transaction
from .models import Patient, UltrasoundExam, UltrasoundImage, FamilyGroup, Appointment
from .forms import PatientForm, UltrasoundExamForm, PatientPasswordChangeForm, PatientProfileForm, PatientUserForm, AppointmentForm, AppointmentUpdateForm, PatientRegistrationForm
from django.db.models import Count, Sum
from django.db.models.functions import ExtractWeek
from django.utils import timezone
from datetime import timedelta
from billing.models import Bill, ServiceType
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.contrib.auth.models import User
from django.contrib.admin.views.decorators import staff_member_required
from django.utils.decorators import method_decorator
from django.contrib.auth.mixins import LoginRequiredMixin, AccessMixin
from django.contrib.auth.decorators import login_required
from django.contrib.auth import REDIRECT_FIELD_NAME
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from django.conf import settings
from functools import wraps

def generate_ultrasound_docx(exam):
    """
    Helper function to generate ultrasound DOCX report.
    Returns the configured Document object.
    """
    # Load the template document from static
    doc = Document(os.path.join(settings.BASE_DIR, 'static', 'docxtemplate.docx'))

    # Function to replace text in document
    def replace_text_in_doc(doc, old_text, new_text):
        for p in doc.paragraphs:
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)

    # Replace date
    replace_text_in_doc(doc, "OCTOBER 09, 2025", exam.exam_date.strftime('%B %d, %Y').upper())

    # Replace examination performed
    replace_text_in_doc(doc, "KUB ULTRASOUND", f"{exam.procedure_type.name} ULTRASOUND".upper())

    # Replace ward
    replace_text_in_doc(doc, "OPD", exam.patient.get_patient_status_display().upper())

    # Replace case number
    for p in doc.paragraphs:
        if "CASE NUMBER" in p.text:
            p.text = p.text.replace("CASE NUMBER	:                                            ", f"CASE NUMBER	: {str(exam.id).zfill(3)}")

    # Name
    for p in doc.paragraphs:
        if "NAME OF PATIENT" in p.text:
            p.text = p.text.replace("NAME OF PATIENT      : ", f"NAME OF PATIENT      : {exam.patient.last_name}, {exam.patient.first_name}")

    # Age
    for p in doc.paragraphs:
        if "AGE" in p.text:
            p.text = p.text.replace("AGE	                	: ", f"AGE	                	: {exam.patient.age or 'N/A'}")

    # Gender
    for p in doc.paragraphs:
        if "GENDER" in p.text:
            p.text = p.text.replace("GENDER	             : ", exam.patient.get_sex_display())

    # Marital status
    for p in doc.paragraphs:
        if "MARITAL STATUS" in p.text:
            p.text = p.text.replace("MARITAL STATUS        :", f"MARITAL STATUS        : {exam.patient.get_marital_status_display() if exam.patient.marital_status else ''}")

    # Requesting physician
    for p in doc.paragraphs:
        if "REQUESTING PHYSICIAN" in p.text:
            p.text = p.text.replace("REQUESTING PHYSICIAN : ", f"REQUESTING PHYSICIAN : {exam.referring_physician or 'N/A'}")

    # Amount paid
    for p in doc.paragraphs:
        if "AMOUNT PAID:" in p.text:
            bill = Bill.objects.filter(items__exam=exam).first()
            amount = bill.total_amount if bill else 0
            p.text = p.text.replace("AMOUNT PAID:", f"AMOUNT PAID: {amount}")

    # Findings
    findings_index = None
    for i, p in enumerate(doc.paragraphs):
        if "ULTRASOUND REPORT:" in p.text:
            findings_index = i
            break
    if findings_index is not None and findings_index + 1 < len(doc.paragraphs):
        doc.paragraphs[findings_index + 1].text = exam.findings or "No specific findings recorded."

    # Impression
    impression_index = None
    for i, p in enumerate(doc.paragraphs):
        if "IMPRESSION :" in p.text:
            impression_index = i
            break
    if impression_index is not None and impression_index + 1 < len(doc.paragraphs):
        doc.paragraphs[impression_index + 1].text = exam.impression or "No impression recorded."

    # Add recommendations if any
    if exam.recommendations or exam.followup_duration or exam.specialist_referral:
        rec_heading = doc.add_paragraph()
        rec_heading_run = rec_heading.add_run('RECOMMENDATIONS')
        rec_heading_run.font.bold = True
        rec_heading_run.font.size = Pt(12)
        rec_para = doc.add_paragraph()
        rec_para.add_run(f"Recommendation: {exam.get_recommendations_display()}")
        if exam.followup_duration:
            rec_para.add_run(f"\nFollow-up Duration: {exam.followup_duration}")
        if exam.specialist_referral:
            rec_para.add_run(f"\nSpecialist Referral: {exam.specialist_referral}")
        doc.add_paragraph()

    # Additional Notes
    if exam.notes:
        notes_heading = doc.add_paragraph()
        notes_heading_run = notes_heading.add_run('ADDITIONAL NOTES')
        notes_heading_run.font.bold = True
        notes_heading_run.font.size = Pt(12)
        notes_para = doc.add_paragraph()
        notes_para.add_run(exam.notes)
        doc.add_paragraph()

    # Technician
    if exam.technician:
        tech_heading = doc.add_paragraph()
        tech_heading_run = tech_heading.add_run('TECHNICIAN')
        tech_heading_run.font.bold = True
        tech_heading_run.font.size = Pt(12)
        tech_para = doc.add_paragraph()
        tech_para.add_run(f"Performed by: {exam.technician}")
        doc.add_paragraph()

    # Images
    if exam.images.exists():
        images_heading = doc.add_paragraph()
        images_heading_run = images_heading.add_run('ULTRASOUND IMAGES')
        images_heading_run.font.bold = True
        images_heading_run.font.size = Pt(12)
        for image in exam.images.all():
            if image.caption:
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"Image: {image.caption}")
                caption_run.font.italic = True
                caption_run.font.size = Pt(10)
            try:
                img_path = image.image.path
                img_para = doc.add_paragraph()
                img_run = img_para.add_run()
                img_run.add_picture(img_path, width=Inches(4.0))
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                original_caption = doc.add_paragraph()
                original_caption_run = original_caption.add_run("Original Ultrasound Image")
                original_caption_run.font.size = Pt(9)
                original_caption_run.font.italic = True
                original_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph("Original image could not be loaded.")
            if image.annotated_image:
                try:
                    annotated_path = image.annotated_image.path
                    annotated_img_para = doc.add_paragraph()
                    annotated_img_run = annotated_img_para.add_run()
                    annotated_img_run.add_picture(annotated_path, width=Inches(4.0))
                    annotated_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    annotated_caption = doc.add_paragraph()
                    annotated_caption_run = annotated_caption.add_run("Annotated Ultrasound Image")
                    annotated_caption_run.font.size = Pt(9)
                    annotated_caption_run.font.italic = True
                    annotated_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph("Annotated image could not be loaded.")
            doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"This report was generated on {timezone.localtime().strftime('%B %d, %Y at %I:%M %p')} by MSRA Ultrasound Clinic Management System.\n\n"
        "IMPORTANT MEDICAL DISCLAIMER: This ultrasound examination report contains preliminary findings and should be interpreted "
        "by a qualified healthcare professional. The final diagnosis and treatment recommendations must be provided by the "
        "attending physician. This report is for medical records purposes only and should not be used as the sole basis "
        "for medical decision-making."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc

def custom_staff_member_required(view_func):
    """
    Custom decorator that requires staff membership and redirects to forbidden page
    instead of Django admin login.
    """
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('forbidden')
        if not request.user.is_staff:
            return redirect('forbidden')
        return view_func(request, *args, **kwargs)
    return wrapper

class CustomStaffRequiredMixin(AccessMixin):
    """
    Custom mixin that requires staff membership and redirects to forbidden page
    instead of Django admin login.
    """
    def dispatch(self, request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('forbidden')
        if not request.user.is_staff:
            return redirect('forbidden')
        return super().dispatch(request, *args, **kwargs)

def require_valid_navigation(view_func):
    """
    Decorator to ensure views are accessed through proper navigation flow.
    Redirects to forbidden page if accessed directly via URL.
    """
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        # Check if request has proper referrer
        referer = request.META.get('HTTP_REFERER', '')
        
        # Allow POST requests (form submissions)
        if request.method == 'POST':
            return view_func(request, *args, **kwargs)
        
        # Allow if referrer is from the same domain and a valid page
        if referer and referer.startswith(request.build_absolute_uri('/')):
            # Check if referrer is from a valid page
            valid_referrer_patterns = [
                '/',
                '/patients/',
                '/patient/',
                '/custom-admin/',
                '/patient-portal/',
                '/patient-appointments/',
                '/staff/appointments/',
                '/home/',
                '/dashboard/',
            ]
            
            for pattern in valid_referrer_patterns:
                if pattern in referer:
                    return view_func(request, *args, **kwargs)
        
        # If no valid referrer, redirect to forbidden page
        return redirect('forbidden')
    
    return wrapper

class PatientListView(CustomStaffRequiredMixin, ListView):
    model = Patient
    template_name = 'patients/patient_list.html'
    context_object_name = 'patients'
    ordering = ['-created_at']
    paginate_by = 10

    def get_queryset(self):
        queryset = super().get_queryset().filter(is_archived=False)
        
        # Handle search
        search_query = self.request.GET.get('search')
        if search_query:
            queryset = queryset.filter(
                models.Q(first_name__icontains=search_query) |
                models.Q(last_name__icontains=search_query) |
                models.Q(contact_number__icontains=search_query) |
                models.Q(email__icontains=search_query) |
                models.Q(id_number__icontains=search_query)
            )
        
        # Handle sex filtering
        sex_filter = self.request.GET.get('sex_filter')
        if sex_filter in ['M', 'F']:
            queryset = queryset.filter(sex=sex_filter)

        # Patient type filter
        patient_type = self.request.GET.get('patient_type')
        if patient_type in dict(Patient.PATIENT_TYPE_CHOICES):
            queryset = queryset.filter(patient_type=patient_type)

        # Patient status filter (in/out)
        patient_status = self.request.GET.get('patient_status')
        if patient_status in dict(Patient.PATIENT_STATUS_CHOICES):
            queryset = queryset.filter(patient_status=patient_status)

        # Location filters
        region = self.request.GET.get('region')
        province = self.request.GET.get('province')
        city = self.request.GET.get('city')
        barangay = self.request.GET.get('barangay')
        if region:
            queryset = queryset.filter(region=region)
        if province:
            queryset = queryset.filter(province=province)
        if city:
            queryset = queryset.filter(city=city)
        if barangay:
            queryset = queryset.filter(barangay=barangay)

        # Created date range filters
        from django.utils.dateparse import parse_date
        created_start = self.request.GET.get('created_start')
        created_end = self.request.GET.get('created_end')
        if created_start:
            start_date = parse_date(created_start)
            if start_date:
                queryset = queryset.filter(created_at__date__gte=start_date)
        if created_end:
            end_date = parse_date(created_end)
            if end_date:
                queryset = queryset.filter(created_at__date__lte=end_date)

        # Age range filters -> convert to birthday range
        age_min = self.request.GET.get('age_min')
        age_max = self.request.GET.get('age_max')
        if age_min or age_max:
            today = timezone.now().date()
            if age_min:
                try:
                    age_min_int = int(age_min)
                    # birthdate <= today - age_min years
                    cutoff = today.replace(year=today.year - age_min_int)
                    queryset = queryset.filter(birthday__lte=cutoff)
                except Exception:
                    pass
            if age_max:
                try:
                    age_max_int = int(age_max)
                    # birthdate >= today - age_max years
                    cutoff = today.replace(year=today.year - age_max_int)
                    queryset = queryset.filter(birthday__gte=cutoff)
                except Exception:
                    pass

        # Last visit date range and has_visits
        queryset = queryset.annotate(last_visit=models.Max('ultrasound_exams__exam_date'))
        last_visit_start = self.request.GET.get('last_visit_start')
        last_visit_end = self.request.GET.get('last_visit_end')
        has_visits = self.request.GET.get('has_visits')  # 'yes' | 'no'
        if last_visit_start:
            start_lv = parse_date(last_visit_start)
            if start_lv:
                queryset = queryset.filter(last_visit__gte=start_lv)
        if last_visit_end:
            end_lv = parse_date(last_visit_end)
            if end_lv:
                queryset = queryset.filter(last_visit__lte=end_lv)
        if has_visits == 'yes':
            queryset = queryset.filter(last_visit__isnull=False)
        elif has_visits == 'no':
            queryset = queryset.filter(last_visit__isnull=True)
        
        # Handle sorting
        sort = self.request.GET.get('sort')
        if sort:
            if sort == 'age_asc':
                queryset = queryset.order_by('age')
            elif sort == 'age_desc':
                queryset = queryset.order_by('-age')
            elif sort == 'visit_asc':
                queryset = queryset.order_by('last_visit')
            elif sort == 'visit_desc':
                queryset = queryset.order_by('-last_visit')
        
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # expose options and current selections
        context['patient_type_choices'] = Patient.PATIENT_TYPE_CHOICES
        context['patient_status_choices'] = Patient.PATIENT_STATUS_CHOICES
        context['current_filters'] = {
            'search': self.request.GET.get('search', ''),
            'sex_filter': self.request.GET.get('sex_filter', ''),
            'patient_type': self.request.GET.get('patient_type', ''),
            'patient_status': self.request.GET.get('patient_status', ''),
            'region': self.request.GET.get('region', ''),
            'province': self.request.GET.get('province', ''),
            'city': self.request.GET.get('city', ''),
            'barangay': self.request.GET.get('barangay', ''),
            'created_start': self.request.GET.get('created_start', ''),
            'created_end': self.request.GET.get('created_end', ''),
            'age_min': self.request.GET.get('age_min', ''),
            'age_max': self.request.GET.get('age_max', ''),
            'last_visit_start': self.request.GET.get('last_visit_start', ''),
            'last_visit_end': self.request.GET.get('last_visit_end', ''),
            'has_visits': self.request.GET.get('has_visits', ''),
            'sort': self.request.GET.get('sort', ''),
        }
        return context

class PatientDetailView(CustomStaffRequiredMixin, DetailView):
    model = Patient
    template_name = 'patients/patient_detail.html'
    context_object_name = 'patient'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Get all ultrasound exams for this patient
        context['exams'] = self.object.ultrasound_exams.all().order_by('-exam_date', '-exam_time')
        return context

class PatientCreateView(CustomStaffRequiredMixin, CreateView):
    model = Patient
    form_class = PatientForm
    template_name = 'patients/patient_form.html'
    success_url = reverse_lazy('patient-list')

    def form_valid(self, form):
        messages.success(self.request, 'Patient record created successfully.')
        return super().form_valid(form)

@method_decorator(custom_staff_member_required, name='dispatch')
@method_decorator(require_valid_navigation, name='dispatch')
class PatientUpdateView(UpdateView):
    model = Patient
    form_class = PatientForm
    template_name = 'patients/patient_form.html'
    
    def get_success_url(self):
        return reverse_lazy('patient-detail', kwargs={'pk': self.object.pk})

    def form_valid(self, form):
        messages.success(self.request, 'Patient record updated successfully.')
        return super().form_valid(form)

    def dispatch(self, request, *args, **kwargs):
        patient = self.get_object()
        if patient.is_archived:
            messages.warning(request, 'Cannot edit an archived patient. Please unarchive first.')
            return redirect('patient-detail', pk=patient.pk)
        return super().dispatch(request, *args, **kwargs)

@method_decorator(custom_staff_member_required, name='dispatch')
@method_decorator(require_valid_navigation, name='dispatch')
class PatientDeleteView(LoginRequiredMixin, DeleteView):
    model = Patient
    template_name = 'patients/patient_confirm_delete.html'

    def get_success_url(self):
        if 'custom-admin' in self.request.path:
            return reverse_lazy('admin_patient_list')
        return reverse_lazy('patient-list')

    def delete(self, request, *args, **kwargs):
        messages.success(request, 'Patient archived successfully.')
        return super().delete(request, *args, **kwargs)

@method_decorator(staff_member_required, name='dispatch')
class ArchivedPatientListView(ListView):
    model = Patient
    template_name = 'patients/archived_patient_list.html'
    context_object_name = 'patients'
    ordering = ['-archived_at']

    def get_queryset(self):
        queryset = super().get_queryset().filter(is_archived=True)
        search_query = self.request.GET.get('search')
        if search_query:
            queryset = queryset.filter(
                models.Q(first_name__icontains=search_query) |
                models.Q(last_name__icontains=search_query)
            )
        return queryset

@custom_staff_member_required
def unarchive_patient(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    patient.is_archived = False
    patient.archived_at = None
    patient.save(update_fields=['is_archived', 'archived_at'])
    messages.success(request, 'Patient restored from archive.')
    return redirect('admin_patient_list')

@method_decorator(staff_member_required, name='dispatch')
@method_decorator(require_valid_navigation, name='dispatch')
class UltrasoundExamCreateView(CreateView):
    model = UltrasoundExam
    form_class = UltrasoundExamForm
    template_name = 'patients/ultrasound_form.html'

    def get_initial(self):
        initial = super().get_initial()
        if 'patient_id' in self.kwargs:
            initial['patient'] = get_object_or_404(Patient, pk=self.kwargs['patient_id'])
        return initial

    def dispatch(self, request, *args, **kwargs):
        if 'patient_id' in kwargs:
            patient = get_object_or_404(Patient, pk=kwargs['patient_id'])
            if patient.is_archived:
                messages.warning(request, 'Cannot add an exam to an archived patient. Please unarchive first.')
                return redirect('patient-detail', pk=patient.pk)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        try:
            # Prevent creating exams for archived patients via form submission
            selected_patient = form.cleaned_data.get('patient')
            if selected_patient and selected_patient.is_archived:
                messages.warning(self.request, 'Cannot add an exam to an archived patient. Please unarchive first.')
                return self.form_invalid(form)
            with transaction.atomic():
                # Save the exam first
                self.object = form.save()

                # Handle multiple image uploads
                files = self.request.FILES.getlist('images[]')
                for file in files:
                    UltrasoundImage.objects.create(
                        exam=self.object,
                        image=file
                    )

                # Send notification to staff about new exam
                from .notification_utils import notify_staff_new_exam
                notify_staff_new_exam(self.object)

                messages.success(self.request, 'Ultrasound examination record created successfully.')
                return super().form_valid(form)
        except Exception as e:
            messages.error(self.request, f'Error saving examination: {str(e)}')
            return self.form_invalid(form)

    def get_success_url(self):
        return reverse_lazy('patient-detail', kwargs={'pk': self.object.patient.pk})

@method_decorator(staff_member_required, name='dispatch')
@method_decorator(require_valid_navigation, name='dispatch')
class UltrasoundExamUpdateView(UpdateView):
    model = UltrasoundExam
    form_class = UltrasoundExamForm
    template_name = 'patients/ultrasound_form.html'

    def form_valid(self, form):
        try:
            with transaction.atomic():
                # Get the old exam instance to check for status changes
                old_exam = UltrasoundExam.objects.get(pk=self.object.pk)

                # Save the exam first
                self.object = form.save()

                # Handle multiple image uploads
                files = self.request.FILES.getlist('images[]')
                for file in files:
                    UltrasoundImage.objects.create(
                        exam=self.object,
                        image=file
                    )

                # Check if status changed to completed and send notification
                if old_exam.status != 'COMPLETED' and self.object.status == 'COMPLETED':
                    from .notification_utils import notify_patient_exam_completed
                    notify_patient_exam_completed(self.object)

                # Send notification to staff about exam update
                from .notification_utils import notify_staff_exam_updated
                notify_staff_exam_updated(self.object)

                messages.success(self.request, 'Ultrasound examination record updated successfully.')
                return super().form_valid(form)
        except Exception as e:
            messages.error(self.request, f'Error updating examination: {str(e)}')
            return self.form_invalid(form)

    def get_success_url(self):
        return reverse_lazy('patient-detail', kwargs={'pk': self.object.patient.pk})

class UltrasoundExamDetailView(CustomStaffRequiredMixin, DetailView):
    model = UltrasoundExam
    template_name = 'patients/ultrasound_detail.html'
    context_object_name = 'exam'

class ImageAnnotationView(CustomStaffRequiredMixin, DetailView):
    model = Patient
    template_name = 'patients/image_annotation.html'
    context_object_name = 'patient'

    def get_object(self, queryset=None):
        if 'image_id' in self.kwargs:
            # If image_id is provided, get the patient through the image
            image = get_object_or_404(UltrasoundImage, id=self.kwargs['image_id'])
            self.specific_image = image
            return image.exam.patient
        return super().get_object(queryset)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if hasattr(self, 'specific_image'):
            # If we're annotating a specific image, only show that exam
            context['exams'] = [self.specific_image.exam]
            context['specific_image_id'] = self.specific_image.id
        else:
            # Otherwise show all exams
            context['exams'] = self.object.ultrasound_exams.all().order_by('-exam_date', '-exam_time')
        
        # Add active procedure types to context
        context['procedure_types'] = ServiceType.objects.filter(is_active=True)
        return context

@custom_staff_member_required
@require_valid_navigation
@require_http_methods(["POST"])
def exam_image_upload(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)
    exam_id = request.POST.get('exam_id')
    image_files = request.FILES.getlist('images[]')
    
    if patient.is_archived:
        messages.error(request, 'Cannot upload images for an archived patient. Please unarchive first.')
        return HttpResponseRedirect(request.META.get('HTTP_REFERER', '/'))

    if not exam_id or not image_files:
        messages.error(request, 'Both examination and images are required.')
        return HttpResponseRedirect(request.META.get('HTTP_REFERER', '/'))
    
    try:
        exam = UltrasoundExam.objects.get(id=exam_id, patient=patient)
        
        # Create UltrasoundImage instances for each uploaded file
        for image_file in image_files:
            UltrasoundImage.objects.create(
                exam=exam,
                image=image_file
            )
        
        messages.success(request, f'{len(image_files)} image(s) uploaded successfully.')
    except UltrasoundExam.DoesNotExist:
        messages.error(request, 'Invalid examination selected.')
    except Exception as e:
        messages.error(request, f'Error uploading images: {str(e)}')
    
    return HttpResponseRedirect(request.META.get('HTTP_REFERER', '/'))

@custom_staff_member_required
@require_http_methods(["POST"])
def delete_ultrasound_image(request, image_id):
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'message': 'Authentication required'}, status=403)
    
    try:
        image = get_object_or_404(UltrasoundImage, pk=image_id)
        image.delete()
        return JsonResponse({'success': True, 'message': 'Image deleted successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=500)

@custom_staff_member_required
def dashboard(request):
    return redirect('admin_dashboard')

@custom_staff_member_required
def home_dashboard(request):
    from django.db.models import Sum, Count
    from datetime import timedelta
    from billing.models import Bill
    from .models import Appointment
    import json
    from collections import defaultdict
    
    # Get search parameter
    patient_search = request.GET.get('patient_search', '')
    
    # Query for patients based on search
    if patient_search:
        searched_patients = Patient.objects.filter(is_archived=False).filter(
            models.Q(first_name__icontains=patient_search) |
            models.Q(last_name__icontains=patient_search) |
            models.Q(contact_number__icontains=patient_search) |
            models.Q(id_number__icontains=patient_search)
        ).order_by('-created_at')[:10]  # Show more results for search
    else:
        searched_patients = None

    # Calculate KPIs
    today = timezone.now().date()
    
    # Basic counts
    total_patients = Patient.objects.filter(is_archived=False).count()
    total_procedures = UltrasoundExam.objects.count()
    
    # Revenue calculations
    total_revenue = Bill.objects.filter(status='PAID').aggregate(Sum('total_amount'))['total_amount__sum'] or 0
    pending_bills = Bill.objects.filter(status='PENDING').count()
    
    # Appointment counts
    pending_appointments = Appointment.objects.filter(status='PENDING').count()
    today_appointments = Appointment.objects.filter(appointment_date=today).count()
    
    # Advanced KPIs
    ninety_days_ago = today - timedelta(days=90)
    six_months_ago = today - timedelta(days=180)
    month_start = today.replace(day=1)

    # Active patients in last 90 days
    active_patients_qs = UltrasoundExam.objects.filter(
        exam_date__gte=ninety_days_ago
    ).values('patient').distinct()
    active_patients_90d = active_patients_qs.count()

    # New patients this month
    new_patients_month = Patient.objects.filter(
        is_archived=False,
        created_at__date__gte=month_start
    ).count()

    # Average procedures per patient (lifetime)
    distinct_patients_with_exam = UltrasoundExam.objects.values('patient').distinct().count()
    avg_procs = (total_procedures / distinct_patients_with_exam) if distinct_patients_with_exam else 0
    avg_procedures_per_patient = f"{avg_procs:.2f}"

    # Returning patient rate (last 6 months)
    recent_exam_counts = (
        UltrasoundExam.objects.filter(exam_date__gte=six_months_ago)
        .values('patient')
        .annotate(num_exams=Count('id'))
    )
    num_recent_unique = recent_exam_counts.count()
    num_returning = sum(1 for r in recent_exam_counts if r['num_exams'] >= 2)
    returning_rate = (num_returning / num_recent_unique * 100) if num_recent_unique else 0
    returning_rate_percent = f"{returning_rate:.1f}"

    # Chart Data
    # Monthly Revenue Data (Last 6 months)
    monthly_revenue_data = []
    monthly_revenue_labels = []
    for i in range(6):
        month_date = today.replace(day=1) - timedelta(days=30*i)
        month_revenue = Bill.objects.filter(
            bill_date__year=month_date.year,
            bill_date__month=month_date.month,
            status__in=['PAID', 'PARTIAL']
        ).aggregate(Sum('total_amount'))['total_amount__sum'] or 0
        monthly_revenue_data.append(float(month_revenue))
        monthly_revenue_labels.append(month_date.strftime('%b %Y'))
    
    monthly_revenue_data.reverse()
    monthly_revenue_labels.reverse()

    # Procedure Distribution Data
    procedure_data = []
    procedure_labels = []
    procedures = UltrasoundExam.objects.values('procedure_type__name').annotate(count=Count('id')).order_by('-count')[:8]
    for proc in procedures:
        procedure_data.append(proc['count'])
        procedure_labels.append(proc['procedure_type__name'] or 'Unknown')

    # Monthly Activity Data (Last 6 months)
    activity_labels = []
    activity_procedures = []
    activity_patients = []
    
    for i in range(6):
        month_date = today.replace(day=1) - timedelta(days=30*i)
        month_procedures = UltrasoundExam.objects.filter(
            exam_date__year=month_date.year,
            exam_date__month=month_date.month
        ).count()
        month_patients = Patient.objects.filter(
            is_archived=False,
            created_at__year=month_date.year,
            created_at__month=month_date.month
        ).count()
        
        activity_labels.append(month_date.strftime('%b'))
        activity_procedures.append(month_procedures)
        activity_patients.append(month_patients)
    
    activity_labels.reverse()
    activity_procedures.reverse()
    activity_patients.reverse()
    
    context = {
        'searched_patients': searched_patients,
        'search_term': patient_search,
        'total_patients': total_patients,
        'total_procedures': total_procedures,
        'total_revenue': f"{total_revenue:,.2f}",
        'pending_bills': pending_bills,
        'active_patients_90d': active_patients_90d,
        'new_patients_month': new_patients_month,
        'avg_procedures_per_patient': avg_procedures_per_patient,
        'returning_rate_percent': returning_rate_percent,
        'pending_appointments': pending_appointments,
        'today_appointments': today_appointments,
        # Chart data
        'monthly_revenue_labels': json.dumps(monthly_revenue_labels),
        'monthly_revenue_data': json.dumps(monthly_revenue_data),
        'procedure_labels': json.dumps(procedure_labels),
        'procedure_data': json.dumps(procedure_data),
        'activity_labels': json.dumps(activity_labels),
        'activity_procedures': json.dumps(activity_procedures),
        'activity_patients': json.dumps(activity_patients),
    }
    
    return render(request, 'home_dashboard.html', context)

@require_valid_navigation
def admin_login(request):
    # Clear all messages
    storage = messages.get_messages(request)
    storage.used = True
    
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        
        if user is not None and user.is_staff:
            # If a staff user is logging into admin, store their original ID
            if request.user.is_authenticated and request.user.is_staff and not request.user.is_superuser:
                request.session['_original_user_id'] = request.user.id
                request.session['_original_is_staff'] = request.user.is_staff
                request.session['_original_is_superuser'] = request.user.is_superuser
            
            login(request, user)
            messages.success(request, f'Welcome back, {user.username}! You have been successfully logged in.')
            next_url = request.POST.get('next') or request.GET.get('next')
            if next_url:
                return redirect(next_url)
            return redirect('admin_dashboard')
        else:
            messages.error(request, 'Invalid credentials or insufficient permissions.')
            return redirect('admin_login')
    
    return render(request, 'admin_login.html')

@custom_staff_member_required
def confirm_family_relationship(request, last_name):
    # This feature has been removed. Redirect back to patient creation.
    messages.info(request, 'Family confirmation step is no longer required.')
    return redirect('patient-create')

@custom_staff_member_required
def family_medical_history(request, family_group_id):
    family_group = get_object_or_404(FamilyGroup, id=family_group_id)
    family_members = family_group.family_members.all()
    
    # Get all exams for all family members
    all_exams = UltrasoundExam.objects.filter(
        patient__in=family_members
    ).select_related('patient', 'procedure_type').order_by('-exam_date')
    
    # Group exams by procedure type
    exams_by_procedure = {}
    for exam in all_exams:
        procedure_name = exam.procedure_type.name
        if procedure_name not in exams_by_procedure:
            exams_by_procedure[procedure_name] = []
        exams_by_procedure[procedure_name].append(exam)
    
    # Calculate statistics and patterns
    procedure_stats = {}
    for procedure_name, exams in exams_by_procedure.items():
        stats = {
            'total_exams': len(exams),
            'members_affected': len(set(exam.patient.id for exam in exams)),
            'common_findings': get_common_findings(exams),
            'latest_exam': {
                'exam_date': exams[0].exam_date.strftime('%Y-%m-%d') if exams else None
            },
            'recommendations': get_recommendation_stats(exams)
        }
        procedure_stats[procedure_name] = stats
    
    context = {
        'family_group': family_group,
        'family_members': family_members,
        'exams_by_procedure': exams_by_procedure,
        'procedure_stats': procedure_stats,
        'total_exams': all_exams.count(),
    }
    
    return render(request, 'patients/family_medical_history.html', context)

def get_common_findings(exams):
    # Extract common words/phrases from findings and impressions
    all_findings = ' '.join([
        f"{exam.findings} {exam.impression}" for exam in exams
    ]).lower()
    
    # You could implement more sophisticated text analysis here
    # For now, we'll just count common medical terms
    common_terms = [
        'mass', 'cyst', 'nodule', 'lesion', 'normal', 'abnormal',
        'inflammation', 'enlarged', 'reduced', 'calcification'
    ]
    
    findings_count = {}
    for term in common_terms:
        count = all_findings.count(term)
        if count > 0:
            findings_count[term] = count
    
    # Sort by frequency
    return dict(sorted(findings_count.items(), key=lambda x: x[1], reverse=True))

def get_recommendation_stats(exams):
    recommendations = {}
    for exam in exams:
        rec = exam.get_recommendations_display()
        recommendations[rec] = recommendations.get(rec, 0) + 1
    return recommendations

@custom_staff_member_required
def generate_report(request, exam_id):
    exam = get_object_or_404(UltrasoundExam, id=exam_id)
    
    # Create a new Word document
    doc = Document()
    
    # Add letterhead
    doc.add_heading('ULTRASOUND EXAMINATION REPORT', 0)
    
    # Add patient information
    doc.add_heading('Patient Information:', level=2)
    doc.add_paragraph(f'Name: {exam.patient.first_name} {exam.patient.last_name}')
    doc.add_paragraph(f'Age: {exam.patient.age}')
    doc.add_paragraph(f'Sex: {exam.patient.get_sex_display()}')
    
    # Add examination details
    doc.add_heading('Examination Details:', level=2)
    doc.add_paragraph(f'Date: {exam.exam_date}')
    doc.add_paragraph(f'Time: {exam.exam_time}')
    doc.add_paragraph(f'Procedure: {exam.procedure_type.name} ULTRASOUND')
    doc.add_paragraph(f'REQUESTING PHYSICIAN: {exam.referring_physician}')
    doc.add_paragraph(f'WARD: {exam.patient.get_patient_status_display()}')
    
    # Add findings
    doc.add_heading('RADIOLOGICAL FINDINGS:', level=2)
    doc.add_paragraph(exam.findings)
    
    # Add impression
    doc.add_heading('IMPRESSION:', level=2)
    doc.add_paragraph(exam.impression)
    
    # Add recommendations
    doc.add_heading('RECOMMENDATIONS:', level=2)
    doc.add_paragraph(f'Follow-up Duration: {exam.followup_duration or "-"}')
    doc.add_paragraph(f'Specialist Referral: {exam.specialist_referral or "-"}')
    
    # Save the document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=ultrasound_report_{exam.id}.docx'
    doc.save(response)
    
    return response

class LandingView(TemplateView):
    template_name = 'landing.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            # Active service types for pricing list
            context['service_types'] = ServiceType.objects.filter(is_active=True).order_by('name')
        except Exception:
            context['service_types'] = []
        return context

def patient_login(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        
        if user is not None and hasattr(user, 'patient'):
            login(request, user)
            messages.success(request, f'Welcome back, {user.username}! You have been successfully logged in.')
            return redirect('patient-portal')
        else:
            messages.error(request, 'Invalid username or password.')
            return render(request, 'patient_login.html')
    
    return render(request, 'patient_login.html')

@login_required
def patient_portal(request):
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    patient = request.user.patient
    context = {
        'patient': patient,
        'recent_exams': patient.ultrasound_exams.order_by('-exam_date', '-exam_time')[:5]
    }
    return render(request, 'patients/patient_portal.html', context)

def patient_logout(request):
    logout(request)
    messages.success(request, 'You have been successfully logged out.')
    return redirect('landing')

def staff_logout(request):
    logout(request)
    messages.success(request, 'You have been successfully logged out.')
    return redirect('staff_login')

def staff_login(request):
    # Clear all messages
    storage = messages.get_messages(request)
    storage.used = True

    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)

        if user is not None and user.is_staff and not user.is_superuser:
            login(request, user)
            # Store this staff user as the last logged in staff user
            request.session['last_staff_user_id'] = user.id
            messages.success(request, f'Welcome back, {user.username}! You have been successfully logged in.')
            next_url = request.POST.get('next') or request.GET.get('next')
            if next_url:
                return redirect(next_url)
            return redirect('home-dashboard')
        else:
            messages.error(request, 'Invalid credentials or insufficient permissions.')
            return redirect('staff_login')

    return render(request, 'staff_login.html')

@login_required
def patient_view_exam(request, exam_id):
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')

    exam = get_object_or_404(UltrasoundExam, id=exam_id)

    # Security check: ensure the patient can only view their own exams
    if exam.patient != request.user.patient:
        messages.error(request, 'Access denied. You can only view your own examinations.')
        return redirect('patient-portal')

    # Get bill information for payment status
    from billing.models import Bill
    bill = Bill.objects.filter(items__exam=exam).first()

    context = {
        'exam': exam,
        'patient': request.user.patient,
        'bill': bill
    }
    return render(request, 'patients/patient_exam_detail.html', context)

@custom_staff_member_required
def download_ultrasound_docx(request, pk):
    exam = get_object_or_404(UltrasoundExam, pk=pk)

    # Load the template document from static
    doc = Document(os.path.join(settings.BASE_DIR, 'static', 'docxtemplate.docx'))

    # Function to replace text in document
    def replace_text_in_doc(doc, old_text, new_text):
        for p in doc.paragraphs:
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)

    # Replace date
    replace_text_in_doc(doc, "OCTOBER 09, 2025", exam.exam_date.strftime('%B %d, %Y').upper())

    # Replace examination performed
    replace_text_in_doc(doc, "KUB ULTRASOUND", f"{exam.procedure_type.name} ULTRASOUND".upper())

    # Replace ward
    replace_text_in_doc(doc, "OPD", exam.patient.get_patient_status_display().upper())

    # Replace case number
    for p in doc.paragraphs:
        if "CASE NUMBER" in p.text:
            p.text = p.text.replace("CASE NUMBER	:                                            ", f"CASE NUMBER	: {str(exam.id).zfill(3)}")

    # Name
    for p in doc.paragraphs:
        if "NAME OF PATIENT" in p.text:
            p.text = p.text.replace("NAME OF PATIENT      : ", f"NAME OF PATIENT      : {exam.patient.last_name}, {exam.patient.first_name}")

    # Age
    for p in doc.paragraphs:
        if "AGE" in p.text:
            p.text = p.text.replace("AGE	                	: ", f"AGE	                	: {exam.patient.age or 'N/A'}")

    # Gender
    for p in doc.paragraphs:
        if "GENDER" in p.text:
            p.text = p.text.replace("GENDER	             : ", exam.patient.get_sex_display())

    # Marital status
    for p in doc.paragraphs:
        if "MARITAL STATUS" in p.text:
            p.text = p.text.replace("MARITAL STATUS        :", f"MARITAL STATUS        : {exam.patient.get_marital_status_display() if exam.patient.marital_status else ''}")

    # Requesting physician
    for p in doc.paragraphs:
        if "REQUESTING PHYSICIAN" in p.text:
            p.text = p.text.replace("REQUESTING PHYSICIAN : ", f"REQUESTING PHYSICIAN : {exam.referring_physician or 'N/A'}")

    # Amount paid
    for p in doc.paragraphs:
        if "AMOUNT PAID:" in p.text:
            bill = Bill.objects.filter(items__exam=exam).first()
            amount = bill.total_amount if bill else 0
            p.text = p.text.replace("AMOUNT PAID:", f"AMOUNT PAID: {amount}")

    # Findings
    findings_index = None
    for i, p in enumerate(doc.paragraphs):
        if "ULTRASOUND REPORT:" in p.text:
            findings_index = i
            break
    if findings_index is not None and findings_index + 1 < len(doc.paragraphs):
        doc.paragraphs[findings_index + 1].text = exam.findings or "No specific findings recorded."

    # Impression
    impression_index = None
    for i, p in enumerate(doc.paragraphs):
        if "IMPRESSION :" in p.text:
            impression_index = i
            break
    if impression_index is not None and impression_index + 1 < len(doc.paragraphs):
        doc.paragraphs[impression_index + 1].text = exam.impression or "No impression recorded."

    # Add recommendations if any
    if exam.recommendations or exam.followup_duration or exam.specialist_referral:
        rec_heading = doc.add_paragraph()
        rec_heading_run = rec_heading.add_run('RECOMMENDATIONS')
        rec_heading_run.font.bold = True
        rec_heading_run.font.size = Pt(12)
        rec_para = doc.add_paragraph()
        rec_para.add_run(f"Recommendation: {exam.get_recommendations_display()}")
        if exam.followup_duration:
            rec_para.add_run(f"\nFollow-up Duration: {exam.followup_duration}")
        if exam.specialist_referral:
            rec_para.add_run(f"\nSpecialist Referral: {exam.specialist_referral}")
        doc.add_paragraph()

    # Additional Notes
    if exam.notes:
        notes_heading = doc.add_paragraph()
        notes_heading_run = notes_heading.add_run('ADDITIONAL NOTES')
        notes_heading_run.font.bold = True
        notes_heading_run.font.size = Pt(12)
        notes_para = doc.add_paragraph()
        notes_para.add_run(exam.notes)
        doc.add_paragraph()

    # Technician
    if exam.technician:
        tech_heading = doc.add_paragraph()
        tech_heading_run = tech_heading.add_run('TECHNICIAN')
        tech_heading_run.font.bold = True
        tech_heading_run.font.size = Pt(12)
        tech_para = doc.add_paragraph()
        tech_para.add_run(f"Performed by: {exam.technician}")
        doc.add_paragraph()

    # Images
    if exam.images.exists():
        images_heading = doc.add_paragraph()
        images_heading_run = images_heading.add_run('ULTRASOUND IMAGES')
        images_heading_run.font.bold = True
        images_heading_run.font.size = Pt(12)
        for image in exam.images.all():
            if image.caption:
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"Image: {image.caption}")
                caption_run.font.italic = True
                caption_run.font.size = Pt(10)
            try:
                img_path = image.image.path
                img_para = doc.add_paragraph()
                img_run = img_para.add_run()
                img_run.add_picture(img_path, width=Inches(4.0))
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                original_caption = doc.add_paragraph()
                original_caption_run = original_caption.add_run("Original Ultrasound Image")
                original_caption_run.font.size = Pt(9)
                original_caption_run.font.italic = True
                original_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph("Original image could not be loaded.")
            if image.annotated_image:
                try:
                    annotated_path = image.annotated_image.path
                    annotated_img_para = doc.add_paragraph()
                    annotated_img_run = annotated_img_para.add_run()
                    annotated_img_run.add_picture(annotated_path, width=Inches(4.0))
                    annotated_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    annotated_caption = doc.add_paragraph()
                    annotated_caption_run = annotated_caption.add_run("Annotated Ultrasound Image")
                    annotated_caption_run.font.size = Pt(9)
                    annotated_caption_run.font.italic = True
                    annotated_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph("Annotated image could not be loaded.")
            doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"This report was generated on {timezone.localtime().strftime('%B %d, %Y at %I:%M %p')} by MSRA Ultrasound Clinic Management System.\n\n"
        "IMPORTANT MEDICAL DISCLAIMER: This ultrasound examination report contains preliminary findings and should be interpreted "
        "by a qualified healthcare professional. The final diagnosis and treatment recommendations must be provided by the "
        "attending physician. This report is for medical records purposes only and should not be used as the sole basis "
        "for medical decision-making."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    patient_name = f"{exam.patient.first_name}_{exam.patient.last_name}"
    procedure = exam.procedure_type.name.replace(" ", "_")
    response['Content-Disposition'] = f'attachment; filename={patient_name}-{procedure}.docx'
    doc.save(response)
    return response
@login_required
def patient_settings(request):
    """Patient settings page."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    patient = request.user.patient
    context = {
        'patient': patient,
    }
    return render(request, 'patients/patient_settings.html', context)

@login_required
def patient_change_password(request):
    """Allow patients to change their password."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    if request.method == 'POST':
        form = PatientPasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)
            messages.success(request, 'Your password has been changed successfully.')
            return redirect('patient-settings')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = PatientPasswordChangeForm(request.user)
    
    context = {
        'form': form,
        'patient': request.user.patient,
    }
    return render(request, 'patients/patient_change_password.html', context)

@login_required
def patient_update_profile(request):
    """Allow patients to update their profile information."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    patient = request.user.patient
    user = request.user
    
    if request.method == 'POST':
        profile_form = PatientProfileForm(request.POST, instance=patient)
        user_form = PatientUserForm(request.POST, instance=user)
        
        if profile_form.is_valid() and user_form.is_valid():
            profile_form.save()
            user_form.save()
            messages.success(request, 'Your profile has been updated successfully.')
            return redirect('patient-settings')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        profile_form = PatientProfileForm(instance=patient)
        user_form = PatientUserForm(instance=user)
    
    context = {
        'profile_form': profile_form,
        'user_form': user_form,
        'patient': patient,
    }
    return render(request, 'patients/patient_update_profile.html', context)

@login_required
def patient_download_exam(request, exam_id):
    """Allow patients to download their examination report."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    exam = get_object_or_404(UltrasoundExam, id=exam_id)
    
    # Security check: ensure the patient can only download their own exams
    if exam.patient != request.user.patient:
        messages.error(request, 'Access denied. You can only download your own examinations.')
        return redirect('patient-portal')
    
     # Load the template document from static
    doc = Document(os.path.join(settings.BASE_DIR, 'static', 'docxtemplate.docx'))

    # Function to replace text in document
    def replace_text_in_doc(doc, old_text, new_text):
        for p in doc.paragraphs:
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)

    # Replace date
    replace_text_in_doc(doc, "OCTOBER 09, 2025", exam.exam_date.strftime('%B %d, %Y').upper())

    # Replace examination performed
    replace_text_in_doc(doc, "KUB ULTRASOUND", f"{exam.procedure_type.name} ULTRASOUND".upper())

    # Replace ward
    replace_text_in_doc(doc, "OPD", exam.patient.get_patient_status_display().upper())

    # Replace case number
    for p in doc.paragraphs:
        if "CASE NUMBER" in p.text:
            p.text = p.text.replace("CASE NUMBER	:                                            ", f"CASE NUMBER	: {str(exam.id).zfill(3)}")

    # Name
    for p in doc.paragraphs:
        if "NAME OF PATIENT" in p.text:
            p.text = p.text.replace("NAME OF PATIENT      : ", f"NAME OF PATIENT      : {exam.patient.last_name}, {exam.patient.first_name}")

    # Age
    for p in doc.paragraphs:
        if "AGE" in p.text:
            p.text = p.text.replace("AGE	                	: ", f"AGE	                	: {exam.patient.age or 'N/A'}")

    # Gender
    for p in doc.paragraphs:
        if "GENDER" in p.text:
            p.text = p.text.replace("GENDER	             : ", exam.patient.get_sex_display())

    # Marital status
    for p in doc.paragraphs:
        if "MARITAL STATUS" in p.text:
            p.text = p.text.replace("MARITAL STATUS        :", f"MARITAL STATUS        : {exam.patient.get_marital_status_display() if exam.patient.marital_status else ''}")

    # Requesting physician
    for p in doc.paragraphs:
        if "REQUESTING PHYSICIAN" in p.text:
            p.text = p.text.replace("REQUESTING PHYSICIAN : ", f"REQUESTING PHYSICIAN : {exam.referring_physician or 'N/A'}")

    # Amount paid
    for p in doc.paragraphs:
        if "AMOUNT PAID:" in p.text:
            bill = Bill.objects.filter(items__exam=exam).first()
            amount = bill.total_amount if bill else 0
            p.text = p.text.replace("AMOUNT PAID:", f"AMOUNT PAID: {amount}")

    # Findings
    findings_index = None
    for i, p in enumerate(doc.paragraphs):
        if "ULTRASOUND REPORT:" in p.text:
            findings_index = i
            break
    if findings_index is not None and findings_index + 1 < len(doc.paragraphs):
        doc.paragraphs[findings_index + 1].text = exam.findings or "No specific findings recorded."

    # Impression
    impression_index = None
    for i, p in enumerate(doc.paragraphs):
        if "IMPRESSION :" in p.text:
            impression_index = i
            break
    if impression_index is not None and impression_index + 1 < len(doc.paragraphs):
        doc.paragraphs[impression_index + 1].text = exam.impression or "No impression recorded."

    # Add recommendations if any
    if exam.recommendations or exam.followup_duration or exam.specialist_referral:
        rec_heading = doc.add_paragraph()
        rec_heading_run = rec_heading.add_run('RECOMMENDATIONS')
        rec_heading_run.font.bold = True
        rec_heading_run.font.size = Pt(12)
        rec_para = doc.add_paragraph()
        rec_para.add_run(f"Recommendation: {exam.get_recommendations_display()}")
        if exam.followup_duration:
            rec_para.add_run(f"\nFollow-up Duration: {exam.followup_duration}")
        if exam.specialist_referral:
            rec_para.add_run(f"\nSpecialist Referral: {exam.specialist_referral}")
        doc.add_paragraph()

    # Additional Notes
    if exam.notes:
        notes_heading = doc.add_paragraph()
        notes_heading_run = notes_heading.add_run('ADDITIONAL NOTES')
        notes_heading_run.font.bold = True
        notes_heading_run.font.size = Pt(12)
        notes_para = doc.add_paragraph()
        notes_para.add_run(exam.notes)
        doc.add_paragraph()

    # Technician
    if exam.technician:
        tech_heading = doc.add_paragraph()
        tech_heading_run = tech_heading.add_run('TECHNICIAN')
        tech_heading_run.font.bold = True
        tech_heading_run.font.size = Pt(12)
        tech_para = doc.add_paragraph()
        tech_para.add_run(f"Performed by: {exam.technician}")
        doc.add_paragraph()

    # Images
    if exam.images.exists():
        images_heading = doc.add_paragraph()
        images_heading_run = images_heading.add_run('ULTRASOUND IMAGES')
        images_heading_run.font.bold = True
        images_heading_run.font.size = Pt(12)
        for image in exam.images.all():
            if image.caption:
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(f"Image: {image.caption}")
                caption_run.font.italic = True
                caption_run.font.size = Pt(10)
            try:
                img_path = image.image.path
                img_para = doc.add_paragraph()
                img_run = img_para.add_run()
                img_run.add_picture(img_path, width=Inches(4.0))
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                original_caption = doc.add_paragraph()
                original_caption_run = original_caption.add_run("Original Ultrasound Image")
                original_caption_run.font.size = Pt(9)
                original_caption_run.font.italic = True
                original_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph("Original image could not be loaded.")
            if image.annotated_image:
                try:
                    annotated_path = image.annotated_image.path
                    annotated_img_para = doc.add_paragraph()
                    annotated_img_run = annotated_img_para.add_run()
                    annotated_img_run.add_picture(annotated_path, width=Inches(4.0))
                    annotated_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    annotated_caption = doc.add_paragraph()
                    annotated_caption_run = annotated_caption.add_run("Annotated Ultrasound Image")
                    annotated_caption_run.font.size = Pt(9)
                    annotated_caption_run.font.italic = True
                    annotated_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph("Annotated image could not be loaded.")
            doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"This report was generated on {timezone.localtime().strftime('%B %d, %Y at %I:%M %p')} by MSRA Ultrasound Clinic Management System.\n\n"
        "IMPORTANT MEDICAL DISCLAIMER: This ultrasound examination report contains preliminary findings and should be interpreted "
        "by a qualified healthcare professional. The final diagnosis and treatment recommendations must be provided by the "
        "attending physician. This report is for medical records purposes only and should not be used as the sole basis "
        "for medical decision-making."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Create the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=ultrasound_report_{exam.patient.last_name}_{exam.exam_date.strftime("%Y%m%d")}.docx'
    
    # Save the document
    doc.save(response)
    
    return response 

@login_required
def patient_appointments(request):
    """Patient appointments page."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    patient = request.user.patient
    appointments = patient.appointments.all().order_by('appointment_date', 'appointment_time')
    
    context = {
        'patient': patient,
        'appointments': appointments,
    }
    return render(request, 'patients/patient_appointments.html', context)

@login_required
def patient_bills(request):
    """Patient bills page: list bills and statuses for the logged-in patient."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')

    patient = request.user.patient
    bills = (
        Bill.objects.filter(patient=patient)
        .order_by('-bill_date')
    )

    context = {
        'patient': patient,
        'bills': bills,
    }
    return render(request, 'patients/patient_bills.html', context)

@login_required
def patient_bill_detail(request, bill_number):
    """Patient bill detail: show items, related exam info, and payment history."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')

    patient = request.user.patient
    bill = get_object_or_404(Bill, bill_number=bill_number, patient=patient)

    bill_items = bill.items.all().select_related('exam', 'service')
    payments = bill.payments.all().order_by('-payment_date')

    total_paid = sum(payment.amount for payment in payments)
    remaining_balance = float(bill.total_amount) - float(total_paid)
    if remaining_balance < 0:
        remaining_balance = 0

    context = {
        'patient': patient,
        'bill': bill,
        'bill_items': bill_items,
        'payments': payments,
        'total_paid': total_paid,
        'remaining_balance': remaining_balance,
    }
    return render(request, 'patients/patient_bill_detail.html', context)

@login_required
def patient_book_appointment(request):
    """Allow patients to book new appointments."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    # Prevent booking if the patient already has a pending appointment
    patient = request.user.patient
    has_pending = Appointment.objects.filter(patient=patient, status='PENDING').exists()
    if has_pending and request.method == 'GET':
        messages.warning(request, 'You already have a pending appointment. Please complete or cancel it before booking a new one.')
        return redirect('patient-appointments')
    
    if request.method == 'POST':
        form = AppointmentForm(request.POST)
        if form.is_valid():
            # Double-check on POST to avoid race conditions / bypass
            if Appointment.objects.filter(patient=patient, status='PENDING').exists():
                messages.warning(request, 'You already have a pending appointment. Please complete or cancel it before booking a new one.')
                return redirect('patient-appointments')

            appointment = form.save(commit=False)
            appointment.patient = patient
            appointment.save()
            
            # Send real-time notification to all staff members
            from .notification_utils import notify_staff_new_appointment
            notify_staff_new_appointment(appointment)
            
            messages.success(request, 'Appointment booked successfully! We will contact you to confirm.')
            return redirect('patient-appointments')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = AppointmentForm()
    
    context = {
        'form': form,
        'patient': patient,
    }
    return render(request, 'patients/patient_book_appointment.html', context)

@login_required
def patient_update_appointment(request, appointment_id):
    """Allow patients to update their appointments."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    appointment = get_object_or_404(Appointment, id=appointment_id)
    
    # Security check: ensure the patient can only update their own appointments
    if appointment.patient != request.user.patient:
        messages.error(request, 'Access denied. You can only update your own appointments.')
        return redirect('patient-appointments')
    
    if request.method == 'POST':
        form = AppointmentUpdateForm(request.POST, instance=appointment)
        if form.is_valid():
            form.save()
            messages.success(request, 'Appointment updated successfully!')
            return redirect('patient-appointments')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = AppointmentUpdateForm(instance=appointment)
    
    context = {
        'form': form,
        'appointment': appointment,
        'patient': request.user.patient,
    }
    return render(request, 'patients/patient_update_appointment.html', context)

@login_required
def patient_cancel_appointment(request, appointment_id):
    """Allow patients to cancel their appointments."""
    if not hasattr(request.user, 'patient'):
        messages.error(request, 'Access denied. This portal is for patients only.')
        return redirect('landing')
    
    appointment = get_object_or_404(Appointment, id=appointment_id)
    
    # Security check: ensure the patient can only cancel their own appointments
    if appointment.patient != request.user.patient:
        messages.error(request, 'Access denied. You can only cancel your own appointments.')
        return redirect('patient-appointments')
    
    if request.method == 'POST':
        appointment.status = 'CANCELLED'
        appointment.save()
        messages.success(request, 'Appointment cancelled successfully.')
        return redirect('patient-appointments')
    
    context = {
        'appointment': appointment,
        'patient': request.user.patient,
    }
    return render(request, 'patients/patient_cancel_appointment.html', context)

@custom_staff_member_required
def patient_list_export_excel(request):
    """Export patient list to Excel format."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from django.http import HttpResponse
    from django.db.models import Q

    # Get filter parameters from request
    search_query = request.GET.get('search', '')
    sex_filter = request.GET.get('sex_filter', '')
    patient_type = request.GET.get('patient_type', '')
    patient_status = request.GET.get('patient_status', '')
    region = request.GET.get('region', '')
    province = request.GET.get('province', '')
    city = request.GET.get('city', '')
    barangay = request.GET.get('barangay', '')
    created_start = request.GET.get('created_start', '')
    created_end = request.GET.get('created_end', '')
    age_min = request.GET.get('age_min', '')
    age_max = request.GET.get('age_max', '')
    last_visit_start = request.GET.get('last_visit_start', '')
    last_visit_end = request.GET.get('last_visit_end', '')
    has_visits = request.GET.get('has_visits', '')
    sort = request.GET.get('sort', '')

    # Build queryset with same filters as PatientListView
    queryset = Patient.objects.filter(is_archived=False)

    # Apply search
    if search_query:
        queryset = queryset.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(contact_number__icontains=search_query) |
            Q(email__icontains=search_query) |
            Q(id_number__icontains=search_query)
        )

    # Apply filters
    if sex_filter in ['M', 'F']:
        queryset = queryset.filter(sex=sex_filter)

    if patient_type in dict(Patient.PATIENT_TYPE_CHOICES):
        queryset = queryset.filter(patient_type=patient_type)

    if patient_status in dict(Patient.PATIENT_STATUS_CHOICES):
        queryset = queryset.filter(patient_status=patient_status)

    # Location filters
    if region:
        queryset = queryset.filter(region=region)
    if province:
        queryset = queryset.filter(province=province)
    if city:
        queryset = queryset.filter(city=city)
    if barangay:
        queryset = queryset.filter(barangay=barangay)

    # Created date range filters
    from django.utils.dateparse import parse_date
    if created_start:
        start_date = parse_date(created_start)
        if start_date:
            queryset = queryset.filter(created_at__date__gte=start_date)
    if created_end:
        end_date = parse_date(created_end)
        if end_date:
            queryset = queryset.filter(created_at__date__lte=end_date)

    # Age range filters
    if age_min or age_max:
        today = timezone.now().date()
        if age_min:
            try:
                age_min_int = int(age_min)
                cutoff = today.replace(year=today.year - age_min_int)
                queryset = queryset.filter(birthday__lte=cutoff)
            except Exception:
                pass
        if age_max:
            try:
                age_max_int = int(age_max)
                cutoff = today.replace(year=today.year - age_max_int)
                queryset = queryset.filter(birthday__gte=cutoff)
            except Exception:
                pass

    # Last visit date range and has_visits
    queryset = queryset.annotate(last_visit=models.Max('ultrasound_exams__exam_date'))
    if last_visit_start:
        start_lv = parse_date(last_visit_start)
        if start_lv:
            queryset = queryset.filter(last_visit__gte=start_lv)
    if last_visit_end:
        end_lv = parse_date(last_visit_end)
        if end_lv:
            queryset = queryset.filter(last_visit__lte=end_lv)
    if has_visits == 'yes':
        queryset = queryset.filter(last_visit__isnull=False)
    elif has_visits == 'no':
        queryset = queryset.filter(last_visit__isnull=True)

    # Apply sorting
    if sort:
        if sort == 'age_asc':
            queryset = queryset.order_by('birthday')
        elif sort == 'age_desc':
            queryset = queryset.order_by('-birthday')
        elif sort == 'visit_asc':
            queryset = queryset.order_by('last_visit')
        elif sort == 'visit_desc':
            queryset = queryset.order_by('-last_visit')
    else:
        queryset = queryset.order_by('-created_at')

    # Create workbook and worksheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Patient List"

    # Define styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    center_align = Alignment(horizontal="center")

    # Headers
    headers = [
        'ID', 'First Name', 'Last Name', 'Age', 'Sex', 'Patient Type', 'Patient Status',
        'Contact Number', 'Email', 'Region', 'Province', 'City', 'Barangay',
        'Street Address', 'ID Number', 'Last Visit', 'Created At'
    ]

    # Write headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align

    # Write data
    for row_num, patient in enumerate(queryset, 2):
        ws.cell(row=row_num, column=1, value=patient.id)
        ws.cell(row=row_num, column=2, value=patient.first_name)
        ws.cell(row=row_num, column=3, value=patient.last_name)
        ws.cell(row=row_num, column=4, value=patient.age or '')
        ws.cell(row=row_num, column=5, value=patient.get_sex_display())
        ws.cell(row=row_num, column=6, value=patient.get_patient_type_display())
        ws.cell(row=row_num, column=7, value=patient.get_patient_status_display())
        ws.cell(row=row_num, column=8, value=patient.contact_number)
        ws.cell(row=row_num, column=9, value=patient.email or '')
        ws.cell(row=row_num, column=10, value=patient.region_name)
        ws.cell(row=row_num, column=11, value=patient.province_name)
        ws.cell(row=row_num, column=12, value=patient.city_name)
        ws.cell(row=row_num, column=13, value=patient.barangay_name)
        ws.cell(row=row_num, column=14, value=patient.street_address)
        ws.cell(row=row_num, column=15, value=patient.id_number or '')
        ws.cell(row=row_num, column=16, value=patient.last_visit.strftime('%Y-%m-%d') if patient.last_visit else '')
        ws.cell(row=row_num, column=17, value=patient.created_at.strftime('%Y-%m-%d %H:%M'))

    # Auto-adjust column widths
    for col_num, column in enumerate(ws.columns, 1):
        max_length = 0
        column_letter = chr(64 + col_num)  # A=1, B=2, etc.

        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass

        adjusted_width = min(max_length + 2, 50)  # Max width of 50
        ws.column_dimensions[column_letter].width = adjusted_width

    # Create response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=patient_list.xlsx'

    # Save workbook to response
    wb.save(response)
    return response

@custom_staff_member_required
def staff_appointments(request):
    """Staff view to manage all appointments."""
    # Get filter parameters
    status_filter = request.GET.get('status', '')
    date_filter = request.GET.get('date', '')
    
    # Base queryset - newest booked first (by creation time)
    appointments = (
        Appointment.objects.select_related('patient')
        .order_by('-created_at', '-appointment_date', '-appointment_time')
    )
    
    # Apply filters
    if status_filter:
        appointments = appointments.filter(status=status_filter)
    
    if date_filter:
        appointments = appointments.filter(appointment_date=date_filter)
    
    # Get statistics
    total_appointments = Appointment.objects.count()
    pending_appointments = Appointment.objects.filter(status='PENDING').count()
    confirmed_appointments = Appointment.objects.filter(status='CONFIRMED').count()
    today_appointments = Appointment.objects.filter(appointment_date=timezone.now().date()).count()
    
    context = {
        'appointments': appointments,
        'total_appointments': total_appointments,
        'pending_appointments': pending_appointments,
        'confirmed_appointments': confirmed_appointments,
        'today_appointments': today_appointments,
        'status_filter': status_filter,
        'date_filter': date_filter,
    }
    return render(request, 'patients/staff_appointments.html', context)

@custom_staff_member_required
def staff_appointment_detail(request, appointment_id):
    """Staff view to see appointment details and manage status."""
    appointment = get_object_or_404(Appointment, id=appointment_id)
    
    if request.method == 'POST':
        new_status = request.POST.get('status')
        if new_status in ['PENDING', 'CONFIRMED', 'CANCELLED', 'COMPLETED']:
            appointment.status = new_status
            appointment.save()
            messages.success(request, f'Appointment status updated to {new_status}.')
            return redirect('staff-appointments')
    
    context = {
        'appointment': appointment,
    }
    return render(request, 'patients/staff_appointment_detail.html', context)

@custom_staff_member_required
def staff_confirm_appointment(request, appointment_id):
    """Staff view to confirm an appointment."""
    appointment = get_object_or_404(Appointment, id=appointment_id)
    
    if request.method == 'POST':
        appointment.status = 'CONFIRMED'
        appointment.save()
        
        # Send real-time notification to patient
        from .notification_utils import notify_patient_appointment_update
        notify_patient_appointment_update(appointment, 'confirmed')
        
        messages.success(request, f'Appointment for {appointment.patient.first_name} {appointment.patient.last_name} has been confirmed.')
        return redirect('staff-appointments')
    
    context = {
        'appointment': appointment,
    }
    return render(request, 'patients/staff_confirm_appointment.html', context)

@custom_staff_member_required
def staff_cancel_appointment(request, appointment_id):
    """Staff view to cancel an appointment."""
    appointment = get_object_or_404(Appointment, id=appointment_id)
    
    if request.method == 'POST':
        appointment.status = 'CANCELLED'
        appointment.save()
        
        # Send real-time notification to patient
        from .notification_utils import notify_patient_appointment_update
        notify_patient_appointment_update(appointment, 'cancelled')
        
        messages.success(request, f'Appointment for {appointment.patient.first_name} {appointment.patient.last_name} has been cancelled.')
        return redirect('staff-appointments')
    
    context = {
        'appointment': appointment,
    }
    return render(request, 'patients/staff_cancel_appointment.html', context)

@custom_staff_member_required
def staff_complete_appointment(request, appointment_id):
    """Staff view to mark an appointment as completed."""
    appointment = get_object_or_404(Appointment, id=appointment_id)
    
    if request.method == 'POST':
        appointment.status = 'COMPLETED'
        appointment.save()
        messages.success(request, f'Appointment for {appointment.patient.first_name} {appointment.patient.last_name} has been marked as completed.')
        return redirect('staff-appointments')
    
    context = {
        'appointment': appointment,
    }
    return render(request, 'patients/staff_complete_appointment.html', context)

@custom_staff_member_required
def elevate_to_admin(request):
    """
    View to temporarily elevate staff user to admin privileges.
    Requires superuser credentials for authentication.
    """
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        # Authenticate superuser credentials
        user = authenticate(request, username=username, password=password)

        if user and user.is_superuser:
            # Store the current user's ID and status before elevating
            if request.user.is_authenticated:
                request.session['_original_user_id'] = request.user.id
                request.session['_original_is_staff'] = request.user.is_staff
                request.session['_original_is_superuser'] = request.user.is_superuser

            # Log in the superuser
            login(request, user)
            
            # Set elevation flag in session
            request.session['elevated_admin'] = True

            # Log the elevation
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"User {request.user.username} elevated to admin privileges")

            messages.success(request, 'Admin privileges granted. You now have access to admin features.')
            return redirect('admin_dashboard')
        else:
            messages.error(request, 'Invalid superuser credentials. Access denied.')

    return render(request, 'admin/admin_login.html')


@custom_staff_member_required
def revert_from_admin(request):
    """
    View to revert from elevated admin privileges back to original staff permissions.
    """
    if request.session.get('elevated_admin', False):
        original_user_id = request.session.get('_original_user_id')
        
        if original_user_id:
            try:
                original_user = User.objects.get(pk=original_user_id)
                # Log out the current admin user
                logout(request)
                # Log in the original staff user
                login(request, original_user)
                
                # Restore original staff/superuser status (though login should handle this)
                # request.user.is_staff = request.session.get('_original_is_staff', False)
                # request.user.is_superuser = request.session.get('_original_is_superuser', False)
                
                # Clear elevation flags
                del request.session['elevated_admin']
                del request.session['_original_user_id']
                del request.session['_original_is_staff']
                del request.session['_original_is_superuser']

                # Log the reversion
                import logging
                logger = logging.getLogger(__name__)
                logger.info(f"User {original_user.username} reverted from admin privileges")

                messages.success(request, 'Returned to staff privileges.')
            except User.DoesNotExist:
                messages.error(request, 'Original staff user not found. Please log in again.')
                logout(request) # Ensure no elevated session remains
                return redirect('staff_login')
        else:
            messages.error(request, 'No original staff user session found. Please log in again.')
            logout(request) # Ensure no elevated session remains
            return redirect('staff_login')
    else:
        messages.info(request, 'No elevated privileges to revert.')

    return redirect('home-dashboard')

def forbidden_page(request):
    """
    Display the forbidden access page.
    This page is shown when users try to access restricted areas or navigate improperly.
    """
    return render(request, 'forbidden.html')

def patient_register(request):
    """Patient registration view."""
    # Check if user is already logged in and has a patient account
    if request.user.is_authenticated and hasattr(request.user, 'patient'):
        messages.info(request, 'You already have a patient account. You are already logged in.')
        return redirect('patient-portal')

    if request.method == 'POST':
        form = PatientRegistrationForm(request.POST)
        if form.is_valid():
            try:
                # Create user account
                user = User.objects.create_user(
                    username=form.cleaned_data['username'],
                    password=form.cleaned_data['password1'],
                    first_name=form.cleaned_data['first_name'],
                    last_name=form.cleaned_data['last_name'],
                    email=form.cleaned_data.get('email', '')
                )

                # Create patient profile
                patient = Patient.objects.create(
                    user=user,
                    first_name=form.cleaned_data['first_name'],
                    last_name=form.cleaned_data['last_name'],
                    birthday=form.cleaned_data['birthday'],
                    sex=form.cleaned_data['sex'],
                    marital_status=form.cleaned_data.get('marital_status'),
                    patient_type=form.cleaned_data['patient_type'],
                    id_number=form.cleaned_data.get('id_number', ''),
                    region=form.cleaned_data['region'],
                    province=form.cleaned_data['province'],
                    city=form.cleaned_data['city'],
                    barangay=form.cleaned_data['barangay'],
                    street_address=form.cleaned_data['street_address'],
                    contact_number=form.cleaned_data['contact_number'],
                    email=form.cleaned_data.get('email', '')
                )

                # Log the user in
                login(request, user)
                messages.success(request, f'Welcome {user.first_name}! Your patient account has been created successfully.')
                return redirect('patient-portal')

            except Exception as e:
                messages.error(request, f'An error occurred during registration: {str(e)}')
                # Clean up user if patient creation failed
                if 'user' in locals():
                    user.delete()
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = PatientRegistrationForm()

    context = {
        'form': form,
    }
    return render(request, 'patient_register.html', context)
